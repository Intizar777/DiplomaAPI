"""
Dashboard export service — generates Excel and Word reports with conditional assessment
for all dashboard roles: GM, Finance, QE, Line Master.
"""
import io
from datetime import date
from decimal import Decimal
from typing import Any, Optional

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

from sqlalchemy.ext.asyncio import AsyncSession
import structlog

from app.schemas.gm_dashboard import OEESummaryResponse, PlanExecutionResponse, DowntimeRankingResponse
from app.schemas.finance_dashboard import (
    SalesBreakdownResponse,
    RevenueTrendResponse,
    TopProductsResponse,
    GroupByType,
    IntervalType,
    SortBy,
)
from app.schemas.qe_dashboard import ParameterTrendsResponse, BatchAnalysisResponse, DefectParetoResponse
from app.schemas.line_master_dashboard import ShiftProgressResponse, ShiftComparisonResponse, DefectSummaryResponse
from app.services.gm_dashboard_service import GroupManagerDashboardService
from app.services.finance_dashboard_service import FinanceManagerDashboardService
from app.services.qe_dashboard_service import QualityEngineerDashboardService
from app.services.line_master_dashboard_service import LineMasterDashboardService

logger = structlog.get_logger()

# ── Status levels ────────────────────────────────────────────────────────────
NORMAL = 0
WARNING = 1
CRITICAL = 2

# ── Excel fill colours (openpyxl ARGB without alpha) ────────────────────────
COLOR_GREEN = "92D050"
COLOR_YELLOW = "FFC000"
COLOR_RED = "FF4444"
COLOR_HEADER = "366092"
COLOR_SUMMARY = "D9D9D9"

STATUS_COLOR = {NORMAL: COLOR_GREEN, WARNING: COLOR_YELLOW, CRITICAL: COLOR_RED}
STATUS_LABEL = {NORMAL: "Норма", WARNING: "Ниже нормы", CRITICAL: "Критически"}


class DashboardExportService:
    """Generates Excel/Word reports for each dashboard role."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    # ── Public export methods ────────────────────────────────────────────────

    async def export_gm(
        self,
        fmt: str,
        period_days: int,
        date_from: date,
        date_to: date,
    ) -> bytes:
        svc = GroupManagerDashboardService(self.db)
        oee = await svc.get_oee_summary(period_days)
        plan = await svc.get_plan_execution(date_from, date_to)
        downtime = await svc.get_downtime_ranking(date_from, date_to)
        logger.info("export_gm", fmt=fmt, period_days=period_days)
        if fmt == "docx":
            return self._word_gm(oee, plan, downtime, date_from, date_to)
        return self._excel_gm(oee, plan, downtime, date_from, date_to)

    async def export_finance(
        self,
        fmt: str,
        date_from: date,
        date_to: date,
    ) -> bytes:
        svc = FinanceManagerDashboardService(self.db)
        breakdown = await svc.get_sales_breakdown(date_from, date_to, GroupByType.channel)
        trend = await svc.get_revenue_trend(date_from, date_to, IntervalType.month)
        products = await svc.get_top_products(date_from, date_to, 10, SortBy.amount)
        logger.info("export_finance", fmt=fmt)
        if fmt == "docx":
            return self._word_finance(breakdown, trend, products, date_from, date_to)
        return self._excel_finance(breakdown, trend, products, date_from, date_to)

    async def export_qe(
        self,
        fmt: str,
        date_from: date,
        date_to: date,
    ) -> bytes:
        svc = QualityEngineerDashboardService(self.db)
        trends = await svc.get_parameter_trends(date_from, date_to)
        batches = await svc.get_batch_analysis(date_from, date_to)
        pareto = await svc.get_defect_pareto(date_from, date_to)
        logger.info("export_qe", fmt=fmt)
        if fmt == "docx":
            return self._word_qe(trends, batches, pareto, date_from, date_to)
        return self._excel_qe(trends, batches, pareto, date_from, date_to)

    async def export_line_master(
        self,
        fmt: str,
        production_date: date,
        date_from: date,
        date_to: date,
    ) -> bytes:
        svc = LineMasterDashboardService(self.db)
        progress = await svc.get_shift_progress(production_date)
        comparison = await svc.get_shift_comparison(date_from, date_to)
        defects = await svc.get_defect_summary(date_from, date_to)
        logger.info("export_line_master", fmt=fmt)
        if fmt == "docx":
            return self._word_line_master(progress, comparison, defects, production_date, date_from, date_to)
        return self._excel_line_master(progress, comparison, defects, production_date, date_from, date_to)

    # ── Assessment helpers ───────────────────────────────────────────────────

    def _assess_oee(self, v: Decimal) -> tuple[str, str, int]:
        """Return (label, color, level) for OEE value (0–100)."""
        if v >= Decimal("75"):
            return STATUS_LABEL[NORMAL], COLOR_GREEN, NORMAL
        if v >= Decimal("65"):
            return STATUS_LABEL[WARNING], COLOR_YELLOW, WARNING
        return STATUS_LABEL[CRITICAL], COLOR_RED, CRITICAL

    def _assess_defect(self, v: Decimal) -> tuple[str, str, int]:
        """Return (label, color, level) for defect rate (0–100)."""
        if v < Decimal("5"):
            return STATUS_LABEL[NORMAL], COLOR_GREEN, NORMAL
        if v <= Decimal("10"):
            return "Повышенный", COLOR_YELLOW, WARNING
        return STATUS_LABEL[CRITICAL], COLOR_RED, CRITICAL

    def _assess_fulfillment(self, v: Decimal) -> tuple[str, str, int]:
        """Return (label, color, level) for plan fulfillment % (0–100)."""
        if v >= Decimal("85"):
            return STATUS_LABEL[NORMAL], COLOR_GREEN, NORMAL
        if v >= Decimal("70"):
            return STATUS_LABEL[WARNING], COLOR_YELLOW, WARNING
        return STATUS_LABEL[CRITICAL], COLOR_RED, CRITICAL

    def _assess_delay(self, hours: Decimal) -> tuple[str, str, int]:
        """Return (label, color, level) for total delay hours."""
        if hours == Decimal("0"):
            return STATUS_LABEL[NORMAL], COLOR_GREEN, NORMAL
        if hours < Decimal("5"):
            return "Незначительный", COLOR_YELLOW, WARNING
        return STATUS_LABEL[CRITICAL], COLOR_RED, CRITICAL

    def _assess_mom_growth(self, v: Optional[Decimal]) -> tuple[str, str, int]:
        """Return (label, color, level) for MoM revenue growth %."""
        if v is None:
            return "Нет данных", COLOR_YELLOW, WARNING
        if v > Decimal("0"):
            return "Рост", COLOR_GREEN, NORMAL
        if v >= Decimal("-10"):
            return "Снижение", COLOR_YELLOW, WARNING
        return "Критическое снижение", COLOR_RED, CRITICAL

    def _assess_share(self, v: Decimal) -> tuple[str, str, int]:
        """Return (label, color, level) for market share % (concentration risk)."""
        if v < Decimal("30"):
            return STATUS_LABEL[NORMAL], COLOR_GREEN, NORMAL
        if v <= Decimal("50"):
            return "Умеренная концентрация", COLOR_YELLOW, WARNING
        return "Высокая концентрация", COLOR_RED, CRITICAL

    # ── Excel shared helpers ─────────────────────────────────────────────────

    def _style_header(self, ws: Any, headers: list[str]) -> None:
        from openpyxl.styles import Font, PatternFill, Alignment
        fill = PatternFill(start_color=COLOR_HEADER, end_color=COLOR_HEADER, fill_type="solid")
        font = Font(bold=True, color="FFFFFF")
        align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.append(headers)
        for cell in ws[ws.max_row]:
            cell.font = font
            cell.fill = fill
            cell.alignment = align
        ws.row_dimensions[ws.max_row].height = 28

    def _color_row(self, ws: Any, row_idx: int, color_hex: str, num_cols: int) -> None:
        from openpyxl.styles import PatternFill
        fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")
        for col in range(1, num_cols + 1):
            ws.cell(row=row_idx, column=col).fill = fill

    def _add_summary(self, ws: Any, values: list) -> None:
        from openpyxl.styles import Font, PatternFill, Alignment
        ws.append(values)
        row = ws.max_row
        fill = PatternFill(start_color=COLOR_SUMMARY, end_color=COLOR_SUMMARY, fill_type="solid")
        font = Font(bold=True)
        for cell in ws[row]:
            cell.fill = fill
            cell.font = font
            cell.alignment = Alignment(horizontal="center")

    def _autofit(self, ws: Any) -> None:
        from openpyxl.utils import get_column_letter
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=8)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 3, 45)

    def _add_title(self, ws: Any, title: str, num_cols: int) -> None:
        from openpyxl.styles import Font, Alignment
        ws.insert_rows(1)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_cols)
        cell = ws.cell(row=1, column=1)
        cell.value = title
        cell.font = Font(bold=True, size=13)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 22

    def _wb_to_bytes(self, wb: Any) -> bytes:
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ── Chart embedding helpers ───────────────────────────────────────────────

    def _add_chart_to_excel(self, ws: Any, png_bytes: bytes, position: str = 'J2') -> None:
        from openpyxl.drawing.image import Image as XLImage
        img = XLImage(io.BytesIO(png_bytes))
        img.width = 600   # pixels (~15 cm at 96 dpi)
        img.height = 360  # pixels (~9 cm at 96 dpi)
        ws.add_image(img, position)

    def _add_chart_to_word(self, doc: Any, png_bytes: bytes, width: float = 5.5) -> None:
        from docx.shared import Inches
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(width))
        doc.add_paragraph()

    def _no_data_chart(self, title: str) -> bytes:
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.text(0.5, 0.5, 'Нет данных', ha='center', va='center',
                fontsize=16, color='grey', transform=ax.transAxes)
        ax.set_title(title, fontsize=13, fontweight='bold', pad=15)
        ax.axis('off')
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    # ── Chart methods: GM ────────────────────────────────────────────────────

    def _chart_gm_oee_bars(self, oee: OEESummaryResponse) -> bytes:
        if not oee.lines:
            return self._no_data_chart('OEE по производственным линиям')
        labels = [l.production_line or f"Линия {i+1}" for i, l in enumerate(oee.lines)]
        values = [float(l.avg_oee) for l in oee.lines]
        colors = [f"#{STATUS_COLOR[self._assess_oee(l.avg_oee)[2]]}" for l in oee.lines]

        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(labels, values, color=colors, edgecolor='white', linewidth=0.5, zorder=2)
        ax.axhline(y=75, color='#366092', linestyle='--', linewidth=2, label='Цель 75%', zorder=3)
        ax.set_title('OEE по производственным линиям', fontsize=13, fontweight='bold', pad=15)
        ax.set_ylabel('OEE %')
        ax.set_ylim(0, 110)
        ax.legend(fontsize=10)
        ax.tick_params(axis='x', rotation=45)
        ax.grid(axis='y', alpha=0.3, zorder=1)
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1,
                    f'{val:.1f}%', ha='center', va='bottom', fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_gm_plan_combo(self, plan: PlanExecutionResponse) -> bytes:
        if not plan.lines:
            return self._no_data_chart('Выполнение производственного плана')
        labels = [l.production_line or f"Линия {i+1}" for i, l in enumerate(plan.lines)]
        targets = [float(l.target_quantity) for l in plan.lines]
        actuals = [float(l.actual_quantity) for l in plan.lines]
        pcts = [float(l.fulfillment_pct) for l in plan.lines]
        x = list(range(len(labels)))

        fig, ax1 = plt.subplots(figsize=(10, 6))
        ax2 = ax1.twinx()
        w = 0.35
        ax1.bar([i - w / 2 for i in x], targets, width=w, label='План (кг)', color='#366092', alpha=0.85)
        ax1.bar([i + w / 2 for i in x], actuals, width=w, label='Факт (кг)', color='#92D050', alpha=0.85)
        ax2.plot(x, pcts, color='#FF6600', linewidth=2, marker='o', markersize=6, label='Выполнение %')
        ax2.axhline(y=85, color='#FFC000', linestyle='--', linewidth=1.5, label='Цель 85%')
        ax2.set_ylabel('Выполнение %', color='#FF6600')
        ax2.set_ylim(0, 130)
        ax1.set_title('Выполнение производственного плана', fontsize=13, fontweight='bold', pad=15)
        ax1.set_ylabel('Количество (кг)')
        ax1.set_xticks(x)
        ax1.set_xticklabels(labels, rotation=45, ha='right')
        ax1.grid(axis='y', alpha=0.3)
        lines1, lbl1 = ax1.get_legend_handles_labels()
        lines2, lbl2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, lbl1 + lbl2, loc='upper left', fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_gm_downtime_pareto(self, downtime: DowntimeRankingResponse) -> bytes:
        if not downtime.lines:
            return self._no_data_chart('Парето: простои по линиям')
        sorted_lines = sorted(downtime.lines, key=lambda l: l.total_delay_hours, reverse=True)
        labels = [l.production_line or f"Линия {i+1}" for i, l in enumerate(sorted_lines)]
        hours = [float(l.total_delay_hours) for l in sorted_lines]
        total = sum(hours)
        cumulative: list[float] = []
        running = 0.0
        for h in hours:
            running += h
            cumulative.append(running / max(total, 0.001) * 100)
        bar_colors = [f"#{STATUS_COLOR[self._assess_delay(l.total_delay_hours)[2]]}" for l in sorted_lines]

        fig, ax1 = plt.subplots(figsize=(10, 6))
        ax2 = ax1.twinx()
        ax1.bar(labels, hours, color=bar_colors, edgecolor='white', alpha=0.85, zorder=2)
        ax2.plot(labels, cumulative, color='#366092', linewidth=2, marker='o', markersize=5, zorder=3)
        ax2.axhline(y=80, color='#FF4444', linestyle='--', linewidth=1.5, label='80%', zorder=3)
        ax2.set_ylabel('Кумулятивный %')
        ax2.set_ylim(0, 110)
        ax1.set_title('Парето: простои по линиям', fontsize=13, fontweight='bold', pad=15)
        ax1.set_ylabel('Часы простоя')
        ax1.tick_params(axis='x', rotation=45)
        ax1.grid(axis='y', alpha=0.3, zorder=1)
        ax2.legend(fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    # ── Chart methods: Finance ───────────────────────────────────────────────

    def _chart_finance_sales_pie(self, breakdown: SalesBreakdownResponse) -> bytes:
        if not breakdown.groups:
            return self._no_data_chart('Разбивка продаж по каналам')
        labels = [g.group_key for g in breakdown.groups]
        sizes = [float(g.amount_share_pct) for g in breakdown.groups]
        colors = [f"#{STATUS_COLOR[self._assess_share(g.amount_share_pct)[2]]}" for g in breakdown.groups]

        fig, ax = plt.subplots(figsize=(9, 7))
        wedges, texts, autotexts = ax.pie(
            sizes, labels=labels, autopct='%1.1f%%',
            colors=colors, startangle=90, pctdistance=0.8,
            textprops={'fontsize': 10},
        )
        for at in autotexts:
            at.set_fontsize(9)
        ax.set_title('Разбивка продаж по каналам', fontsize=13, fontweight='bold', pad=15)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_finance_trend_line(self, trend: RevenueTrendResponse) -> bytes:
        if not trend.data:
            return self._no_data_chart('Тренд выручки')
        dates = [p.trend_date.strftime('%d.%m') for p in trend.data]
        amounts = [float(p.total_amount) for p in trend.data]
        growths = [float(p.mom_growth_pct) if p.mom_growth_pct is not None else 0.0 for p in trend.data]
        x = list(range(len(dates)))

        fig, ax1 = plt.subplots(figsize=(10, 6))
        ax2 = ax1.twinx()
        ax1.plot(x, amounts, color='#366092', linewidth=2.5, marker='o', markersize=5, label='Выручка')
        ax1.fill_between(x, amounts, alpha=0.15, color='#366092')
        bar_colors = ['#92D050' if g >= 0 else '#FF4444' for g in growths]
        ax2.bar(x, growths, color=bar_colors, alpha=0.6, label='Рост MoM %', width=0.4)
        ax2.axhline(y=0, color='grey', linestyle='-', linewidth=0.8)
        ax2.set_ylabel('Рост MoM %')
        ax1.set_title('Тренд выручки', fontsize=13, fontweight='bold', pad=15)
        ax1.set_ylabel('Выручка')
        ax1.set_xticks(x)
        ax1.set_xticklabels(dates, rotation=45, ha='right')
        ax1.grid(axis='y', alpha=0.3)
        lines1, lbl1 = ax1.get_legend_handles_labels()
        lines2, lbl2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, lbl1 + lbl2, loc='upper left', fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_finance_top_barh(self, products: TopProductsResponse) -> bytes:
        if not products.products:
            return self._no_data_chart('Топ продукты по выручке')
        top = products.products[:10]
        top_rev = list(reversed(top))
        labels = [p.product_name[:30] for p in top_rev]
        values = [float(p.total_amount) for p in top_rev]
        colors = [f"#{STATUS_COLOR[self._assess_share(p.amount_share_pct)[2]]}" for p in top_rev]

        fig, ax = plt.subplots(figsize=(10, max(6, len(labels) * 0.45 + 1)))
        bars = ax.barh(labels, values, color=colors, edgecolor='white', alpha=0.85)
        ax.set_title('Топ продукты по выручке', fontsize=13, fontweight='bold', pad=15)
        ax.set_xlabel('Выручка')
        ax.grid(axis='x', alpha=0.3)
        for bar, val in zip(bars, values):
            ax.text(bar.get_width() * 1.01, bar.get_y() + bar.get_height() / 2,
                    f'{val:,.0f}', va='center', fontsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    # ── Chart methods: QE ────────────────────────────────────────────────────

    def _chart_qe_params_bars(self, trends: ParameterTrendsResponse) -> bytes:
        if not trends.parameters:
            return self._no_data_chart('Тренды параметров качества')
        labels = [p.parameter_name[:25] for p in trends.parameters]
        values = [float(p.overall_out_of_spec_pct) for p in trends.parameters]
        colors = [f"#{STATUS_COLOR[self._assess_defect(p.overall_out_of_spec_pct)[2]]}" for p in trends.parameters]
        x = list(range(len(labels)))

        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(x, values, color=colors, edgecolor='white', alpha=0.85, zorder=2)
        ax.axhline(y=5, color='#FFC000', linestyle='--', linewidth=2, label='Порог 5%', zorder=3)
        ax.axhline(y=10, color='#FF4444', linestyle='--', linewidth=2, label='Порог 10%', zorder=3)
        ax.set_title('Тренды параметров качества (% вне допуска)', fontsize=12, fontweight='bold', pad=15)
        ax.set_ylabel('% вне допуска')
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=45, ha='right')
        ax.legend(fontsize=10)
        ax.grid(axis='y', alpha=0.3, zorder=1)
        for i, val in enumerate(values):
            ax.text(i, val + 0.2, f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_qe_batches_timeline(self, batches: BatchAnalysisResponse) -> bytes:
        if not batches.lots:
            return self._no_data_chart('Анализ партий: нарушения по партиям')
        sorted_lots = sorted(
            batches.lots,
            key=lambda l: (l.production_date or date.min, l.lot_number),
        )
        labels = [l.lot_number for l in sorted_lots]
        values = [l.fail_count for l in sorted_lots]
        colors = [f"#{STATUS_COLOR[CRITICAL if l.fail_count > 3 else WARNING]}" for l in sorted_lots]
        x = list(range(len(labels)))

        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(x, values, color=colors, edgecolor='white', alpha=0.85, zorder=2)
        ax.plot(x, values, color='#366092', linewidth=1.5, marker='o', markersize=4, zorder=3)
        ax.axhline(y=3, color='#FFC000', linestyle='--', linewidth=1.5, label='Порог 3', zorder=3)
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=8)
        ax.set_title('Анализ партий: нарушения по партиям', fontsize=12, fontweight='bold', pad=15)
        ax.set_ylabel('Количество нарушений')
        ax.legend(fontsize=10)
        ax.grid(axis='y', alpha=0.3, zorder=1)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_qe_pareto(self, pareto: DefectParetoResponse) -> bytes:
        if not pareto.parameters:
            return self._no_data_chart('Pareto-анализ дефектов')
        labels = [p.parameter_name[:20] for p in pareto.parameters]
        counts = [p.defect_count for p in pareto.parameters]
        cumulatives = [float(p.cumulative_pct) for p in pareto.parameters]
        colors = [f"#{STATUS_COLOR[self._assess_defect(p.defect_pct)[2]]}" for p in pareto.parameters]
        x = list(range(len(labels)))

        fig, ax1 = plt.subplots(figsize=(10, 6))
        ax2 = ax1.twinx()
        ax1.bar(x, counts, color=colors, edgecolor='white', alpha=0.85, zorder=2)
        ax2.plot(x, cumulatives, color='#366092', linewidth=2, marker='o', markersize=5, zorder=3)
        ax2.axhline(y=80, color='#FF4444', linestyle='--', linewidth=1.5, label='80%', zorder=3)
        ax2.set_ylabel('Кумулятивный %')
        ax2.set_ylim(0, 110)
        ax1.set_title('Pareto-анализ дефектов', fontsize=13, fontweight='bold', pad=15)
        ax1.set_ylabel('Количество дефектов')
        ax1.set_xticks(x)
        ax1.set_xticklabels(labels, rotation=45, ha='right', fontsize=9)
        ax1.grid(axis='y', alpha=0.3, zorder=1)
        ax2.legend(fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    # ── Chart methods: Line Master ───────────────────────────────────────────

    def _chart_lm_shift_combo(self, progress: ShiftProgressResponse) -> bytes:
        if not progress.shifts:
            return self._no_data_chart('Прогресс смен: выпуск и дефектность')
        labels = [s.shift for s in progress.shifts]
        qtys = [float(s.total_quantity) for s in progress.shifts]
        rates = [float(s.defect_rate) for s in progress.shifts]
        bar_colors = [f"#{STATUS_COLOR[self._assess_defect(s.defect_rate)[2]]}" for s in progress.shifts]
        x = list(range(len(labels)))

        fig, ax1 = plt.subplots(figsize=(10, 6))
        ax2 = ax1.twinx()
        ax1.bar(x, qtys, color=bar_colors, edgecolor='white', alpha=0.85, label='Кол-во (кг)', zorder=2)
        ax2.plot(x, rates, color='#FF6600', linewidth=2.5, marker='o', markersize=7, label='Дефектность %', zorder=3)
        ax2.axhline(y=5, color='#FFC000', linestyle='--', linewidth=1.5, label='Порог 5%')
        ax2.set_ylabel('Дефектность %', color='#FF6600')
        ax2.set_ylim(0, max(max(rates) * 1.5, 15) if rates else 15)
        ax1.set_title('Прогресс смен: выпуск и дефектность', fontsize=12, fontweight='bold', pad=15)
        ax1.set_ylabel('Количество (кг)')
        ax1.set_xticks(x)
        ax1.set_xticklabels(labels)
        ax1.grid(axis='y', alpha=0.3, zorder=1)
        lines1, lbl1 = ax1.get_legend_handles_labels()
        lines2, lbl2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, lbl1 + lbl2, loc='upper right', fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_lm_comparison_line(self, comparison: ShiftComparisonResponse) -> bytes:
        if not comparison.shifts:
            return self._no_data_chart('Динамика выпуска по сменам')
        labels = [f"{p.date.strftime('%d.%m')}\n{p.shift or 'День'}" for p in comparison.shifts]
        qtys = [float(p.total_quantity) for p in comparison.shifts]
        x = list(range(len(labels)))

        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(x, qtys, color='#366092', linewidth=2.5, marker='o', markersize=6, zorder=3)
        ax.fill_between(x, qtys, alpha=0.15, color='#366092')
        ax.set_title('Динамика выпуска по сменам', fontsize=13, fontweight='bold', pad=15)
        ax.set_ylabel('Количество (кг)')
        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=8)
        ax.grid(axis='y', alpha=0.3, zorder=1)
        max_qty = max(qtys) if qtys else 1
        for i, val in enumerate(qtys):
            ax.text(i, val + max_qty * 0.01, f'{val:.0f}', ha='center', va='bottom', fontsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_lm_defects_barh(self, defects: DefectSummaryResponse) -> bytes:
        if not defects.items:
            return self._no_data_chart('Сводка дефектов по параметрам')
        items_rev = list(reversed(defects.items))
        labels = [i.parameter_name[:30] for i in items_rev]
        values = [float(i.fail_rate) for i in items_rev]
        colors = [f"#{STATUS_COLOR[self._assess_defect(i.fail_rate)[2]]}" for i in items_rev]

        fig, ax = plt.subplots(figsize=(10, max(6, len(labels) * 0.45 + 1)))
        bars = ax.barh(labels, values, color=colors, edgecolor='white', alpha=0.85, zorder=2)
        ax.axvline(x=5, color='#FFC000', linestyle='--', linewidth=2, label='Порог 5%', zorder=3)
        ax.axvline(x=10, color='#FF4444', linestyle='--', linewidth=2, label='Порог 10%', zorder=3)
        ax.set_title('Сводка дефектов по параметрам', fontsize=13, fontweight='bold', pad=15)
        ax.set_xlabel('% провалов')
        ax.legend(fontsize=10)
        ax.grid(axis='x', alpha=0.3, zorder=1)
        for bar, val in zip(bars, values):
            ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                    f'{val:.1f}%', va='center', fontsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    # ── Excel: GM ────────────────────────────────────────────────────────────

    def _excel_gm(
        self,
        oee: OEESummaryResponse,
        plan: PlanExecutionResponse,
        downtime: DowntimeRankingResponse,
        date_from: date,
        date_to: date,
    ) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()

        # Sheet 1 — OEE
        ws1 = wb.active
        ws1.title = "OEE по линиям"
        headers1 = ["Линия", "OEE%", "Цель%", "Отклонение", "Заказов выполнено", "Всего заказов", "Дефекты%", "Вывод"]
        self._style_header(ws1, headers1)

        oee_levels: list[int] = []
        for item in oee.lines:
            label, color, level = self._assess_oee(item.avg_oee)
            oee_levels.append(level)
            ws1.append([
                item.production_line or "Все линии",
                float(item.avg_oee),
                75.0,
                float(item.vs_target_pct),
                item.completed_orders,
                item.total_orders,
                float(item.avg_defect_rate),
                label,
            ])
            self._color_row(ws1, ws1.max_row, color, len(headers1))

        avg_oee = sum(float(i.avg_oee) for i in oee.lines) / max(len(oee.lines), 1)
        self._add_summary(ws1, ["ИТОГО", round(avg_oee, 1), 75.0, round(avg_oee - 75.0, 1),
                                 sum(i.completed_orders for i in oee.lines),
                                 sum(i.total_orders for i in oee.lines), "", ""])
        self._add_title(ws1, f"OEE по линиям | {date_from} — {date_to}", len(headers1))
        self._autofit(ws1)
        self._add_chart_to_excel(ws1, self._chart_gm_oee_bars(oee), 'J2')

        # Sheet 2 — Plan execution
        ws2 = wb.create_sheet("Выполнение плана")
        headers2 = ["Линия", "Цель (кг)", "Факт (кг)", "Выполнение%", "В работе", "Завершено", "Отменено", "Вывод"]
        self._style_header(ws2, headers2)

        plan_levels: list[int] = []
        for plan_line in plan.lines:
            label, color, level = self._assess_fulfillment(plan_line.fulfillment_pct)
            plan_levels.append(level)
            ws2.append([
                plan_line.production_line or "Без линии",
                float(plan_line.target_quantity),
                float(plan_line.actual_quantity),
                float(plan_line.fulfillment_pct),
                plan_line.orders_in_progress,
                plan_line.orders_completed,
                plan_line.orders_cancelled,
                label,
            ])
            self._color_row(ws2, ws2.max_row, color, len(headers2))

        self._add_summary(ws2, ["ИТОГО", float(plan.total_target), float(plan.total_actual),
                                  float(plan.overall_fulfillment_pct), "", "", "", ""])
        self._add_title(ws2, f"Выполнение плана | {date_from} — {date_to}", len(headers2))
        self._autofit(ws2)
        self._add_chart_to_excel(ws2, self._chart_gm_plan_combo(plan), 'J2')

        # Sheet 3 — Downtime
        ws3 = wb.create_sheet("Простои")
        headers3 = ["Линия", "Часы простоя", "Задержанных заказов", "Ср. задержка (ч)", "Всего завершено", "Задержано%", "Вывод"]
        self._style_header(ws3, headers3)

        delay_levels: list[int] = []
        for delay_line in downtime.lines:
            label, color, level = self._assess_delay(delay_line.total_delay_hours)
            delay_levels.append(level)
            ws3.append([
                delay_line.production_line or "Без линии",
                float(delay_line.total_delay_hours),
                delay_line.delayed_orders,
                float(delay_line.avg_delay_hours),
                delay_line.total_completed,
                float(delay_line.delay_pct),
                label,
            ])
            self._color_row(ws3, ws3.max_row, color, len(headers3))

        self._add_summary(ws3, ["ИТОГО", float(downtime.total_delay_hours), downtime.total_delayed_orders,
                                  "", "", "", ""])
        self._add_title(ws3, f"Простои и задержки | {date_from} — {date_to}", len(headers3))
        self._autofit(ws3)
        self._add_chart_to_excel(ws3, self._chart_gm_downtime_pareto(downtime), 'J2')

        return self._wb_to_bytes(wb)

    # ── Excel: Finance ───────────────────────────────────────────────────────

    def _excel_finance(
        self,
        breakdown: SalesBreakdownResponse,
        trend: RevenueTrendResponse,
        products: TopProductsResponse,
        date_from: date,
        date_to: date,
    ) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()

        # Sheet 1 — Sales breakdown
        ws1 = wb.active
        ws1.title = "Разбивка продаж"
        headers1 = ["Группа", "Выручка", "Кол-во", "Продаж", "Ср. чек", "Доля%", "Вывод"]
        self._style_header(ws1, headers1)

        for item in breakdown.groups:
            label, color, _ = self._assess_share(item.amount_share_pct)
            ws1.append([
                item.group_key,
                float(item.total_amount),
                float(item.total_quantity),
                item.sales_count,
                float(item.avg_order_value),
                float(item.amount_share_pct),
                label,
            ])
            self._color_row(ws1, ws1.max_row, color, len(headers1))

        self._add_summary(ws1, ["ИТОГО", float(breakdown.total_amount), float(breakdown.total_quantity),
                                  "", "", 100.0, ""])
        self._add_title(ws1, f"Разбивка продаж | {date_from} — {date_to}", len(headers1))
        self._autofit(ws1)
        self._add_chart_to_excel(ws1, self._chart_finance_sales_pie(breakdown), 'J2')

        # Sheet 2 — Revenue trend
        ws2 = wb.create_sheet("Тренд выручки")
        headers2 = ["Дата", "Выручка", "Кол-во", "Заказов", "Рост MoM%", "Вывод"]
        self._style_header(ws2, headers2)

        for point in trend.data:
            label, color, _ = self._assess_mom_growth(point.mom_growth_pct)
            ws2.append([
                point.trend_date.strftime("%d.%m.%Y"),
                float(point.total_amount),
                float(point.total_quantity),
                point.order_count,
                float(point.mom_growth_pct) if point.mom_growth_pct is not None else None,
                label,
            ])
            self._color_row(ws2, ws2.max_row, color, len(headers2))

        self._add_title(ws2, f"Тренд выручки | {date_from} — {date_to}", len(headers2))
        self._autofit(ws2)
        self._add_chart_to_excel(ws2, self._chart_finance_trend_line(trend), 'J2')

        # Sheet 3 — Top products
        ws3 = wb.create_sheet("Топ продукты")
        headers3 = ["Ранг", "Продукт", "Выручка", "Кол-во", "Продаж", "Доля%", "Вывод"]
        self._style_header(ws3, headers3)

        for prod in products.products:
            label, color, _ = self._assess_share(prod.amount_share_pct)
            ws3.append([
                prod.rank,
                prod.product_name,
                float(prod.total_amount),
                float(prod.total_quantity),
                prod.sales_count,
                float(prod.amount_share_pct),
                label,
            ])
            self._color_row(ws3, ws3.max_row, color, len(headers3))

        self._add_summary(ws3, ["", "ИТОГО", float(products.total_amount), "", "", 100.0, ""])
        self._add_title(ws3, f"Топ продукты | {date_from} — {date_to}", len(headers3))
        self._autofit(ws3)
        self._add_chart_to_excel(ws3, self._chart_finance_top_barh(products), 'J2')

        return self._wb_to_bytes(wb)

    # ── Excel: QE ────────────────────────────────────────────────────────────

    def _excel_qe(
        self,
        trends: ParameterTrendsResponse,
        batches: BatchAnalysisResponse,
        pareto: DefectParetoResponse,
        date_from: date,
        date_to: date,
    ) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()

        # Sheet 1 — Parameter trends
        ws1 = wb.active
        ws1.title = "Тренды параметров"
        headers1 = ["Параметр", "Тестов", "Вне допуска", "% вне допуска", "Вывод"]
        self._style_header(ws1, headers1)

        for item in trends.parameters:
            label, color, _ = self._assess_defect(item.overall_out_of_spec_pct)
            ws1.append([
                item.parameter_name,
                item.total_tests,
                item.total_out_of_spec,
                float(item.overall_out_of_spec_pct),
                label,
            ])
            self._color_row(ws1, ws1.max_row, color, len(headers1))

        total_tests = sum(i.total_tests for i in trends.parameters)
        total_oos = sum(i.total_out_of_spec for i in trends.parameters)
        avg_oos = round(total_oos / max(total_tests, 1) * 100, 2)
        self._add_summary(ws1, ["ИТОГО", total_tests, total_oos, avg_oos, ""])
        self._add_title(ws1, f"Тренды параметров качества | {date_from} — {date_to}", len(headers1))
        self._autofit(ws1)
        self._add_chart_to_excel(ws1, self._chart_qe_params_bars(trends), 'G2')

        # Sheet 2 — Batch analysis
        ws2 = wb.create_sheet("Анализ партий")
        headers2 = ["Партия", "Продукт", "Дата", "Смена", "Нарушений", "Вывод"]
        self._style_header(ws2, headers2)

        for lot in batches.lots:
            level = CRITICAL if lot.fail_count > 3 else WARNING
            label = STATUS_LABEL[level]
            color = STATUS_COLOR[level]
            ws2.append([
                lot.lot_number,
                lot.product_name or "",
                lot.production_date.strftime("%d.%m.%Y") if lot.production_date else "",
                lot.shift or "",
                lot.fail_count,
                label,
            ])
            self._color_row(ws2, ws2.max_row, color, len(headers2))

        self._add_summary(ws2, ["ИТОГО", "", "", "", batches.lot_count, ""])
        self._add_title(ws2, f"Анализ партий с отклонениями | {date_from} — {date_to}", len(headers2))
        self._autofit(ws2)
        self._add_chart_to_excel(ws2, self._chart_qe_batches_timeline(batches), 'H2')

        # Sheet 3 — Pareto
        ws3 = wb.create_sheet("Pareto дефектов")
        headers3 = ["Параметр", "Дефектов", "Тестов", "% дефектов", "Кумул.%", "Вывод"]
        self._style_header(ws3, headers3)

        for pitem in pareto.parameters:
            label, color, _ = self._assess_defect(pitem.defect_pct)
            ws3.append([
                pitem.parameter_name,
                pitem.defect_count,
                pitem.total_tests,
                float(pitem.defect_pct),
                float(pitem.cumulative_pct),
                label,
            ])
            self._color_row(ws3, ws3.max_row, color, len(headers3))

        self._add_summary(ws3, ["ИТОГО", pareto.total_defects, "", "", 100.0, ""])
        self._add_title(ws3, f"Pareto дефектов | {date_from} — {date_to}", len(headers3))
        self._autofit(ws3)
        self._add_chart_to_excel(ws3, self._chart_qe_pareto(pareto), 'H2')

        return self._wb_to_bytes(wb)

    # ── Excel: Line Master ───────────────────────────────────────────────────

    def _excel_line_master(
        self,
        progress: ShiftProgressResponse,
        comparison: ShiftComparisonResponse,
        defects: DefectSummaryResponse,
        production_date: date,
        date_from: date,
        date_to: date,
    ) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()

        # Sheet 1 — Shift progress
        ws1 = wb.active
        ws1.title = "Прогресс смен"
        headers1 = ["Смена", "Лотов", "Кол-во (кг)", "Одобрено", "Дефектов", "Дефектность%", "Вывод"]
        self._style_header(ws1, headers1)

        for shift in progress.shifts:
            label, color, _ = self._assess_defect(shift.defect_rate)
            ws1.append([
                shift.shift,
                shift.lot_count,
                float(shift.total_quantity),
                shift.approved_count,
                shift.defect_count,
                float(shift.defect_rate),
                label,
            ])
            self._color_row(ws1, ws1.max_row, color, len(headers1))

        self._add_summary(ws1, ["ИТОГО", progress.total_lots, float(progress.total_quantity), "", "", "", ""])
        self._add_title(ws1, f"Прогресс смен | {production_date}", len(headers1))
        self._autofit(ws1)
        self._add_chart_to_excel(ws1, self._chart_lm_shift_combo(progress), 'I2')

        # Sheet 2 — Shift comparison
        ws2 = wb.create_sheet("Сравнение смен")
        headers2 = ["Дата", "Смена", "Лотов", "Кол-во (кг)", "Дефектов", "Вывод"]
        self._style_header(ws2, headers2)

        for period in comparison.shifts:
            total_qty = float(period.total_quantity)
            defect_pct = Decimal(period.defect_count) / max(Decimal(period.lot_count), Decimal("1")) * 100
            label, color, _ = self._assess_defect(defect_pct)
            ws2.append([
                period.date.strftime("%d.%m.%Y"),
                period.shift or "День",
                period.lot_count,
                total_qty,
                period.defect_count,
                label,
            ])
            self._color_row(ws2, ws2.max_row, color, len(headers2))

        self._add_title(ws2, f"Сравнение смен | {date_from} — {date_to}", len(headers2))
        self._autofit(ws2)
        self._add_chart_to_excel(ws2, self._chart_lm_comparison_line(comparison), 'H2')

        # Sheet 3 — Defect summary
        ws3 = wb.create_sheet("Сводка дефектов")
        headers3 = ["Параметр", "Тестов", "Провалено", "Провал%", "Вывод"]
        self._style_header(ws3, headers3)

        for item in defects.items:
            label, color, _ = self._assess_defect(item.fail_rate)
            ws3.append([
                item.parameter_name,
                item.total_tests,
                item.failed_tests,
                float(item.fail_rate),
                label,
            ])
            self._color_row(ws3, ws3.max_row, color, len(headers3))

        self._add_summary(ws3, ["ИТОГО", "", defects.total_defects, "", ""])
        self._add_title(ws3, f"Сводка дефектов | {date_from} — {date_to}", len(headers3))
        self._autofit(ws3)
        self._add_chart_to_excel(ws3, self._chart_lm_defects_barh(defects), 'G2')

        return self._wb_to_bytes(wb)

    # ── Word shared helpers ──────────────────────────────────────────────────

    def _word_header(self, doc: Any, role_name: str, date_from: date, date_to: date) -> None:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title = doc.add_heading("АНАЛИТИЧЕСКИЙ ОТЧЁТ", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_heading(role_name, 1)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run("Период анализа: ").bold = True
        p.add_run(f"{date_from.strftime('%d.%m.%Y')} — {date_to.strftime('%d.%m.%Y')}")

        p2 = doc.add_paragraph()
        p2.add_run("Дата составления: ").bold = True
        p2.add_run(date.today().strftime("%d.%m.%Y"))
        doc.add_paragraph()

    def _word_section(
        self,
        doc: Any,
        title: str,
        bullets: list[str],
        conclusion: str,
        recommendation: str,
        chart_bytes: Optional[bytes] = None,
    ) -> None:
        doc.add_heading(title, 2)
        if chart_bytes:
            self._add_chart_to_word(doc, chart_bytes)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")
        p = doc.add_paragraph()
        p.add_run("Вывод: ").bold = True
        p.add_run(conclusion)
        p2 = doc.add_paragraph()
        p2.add_run("Рекомендации: ").bold = True
        p2.add_run(recommendation)
        doc.add_paragraph()

    def _word_to_bytes(self, doc: Any) -> bytes:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def _worst(self, levels: list[int]) -> int:
        return max(levels) if levels else NORMAL

    # ── Word: GM ─────────────────────────────────────────────────────────────

    def _word_gm(
        self,
        oee: OEESummaryResponse,
        plan: PlanExecutionResponse,
        downtime: DowntimeRankingResponse,
        date_from: date,
        date_to: date,
    ) -> bytes:
        from docx import Document

        doc = Document()
        self._word_header(doc, "Производственный директор (Group Manager)", date_from, date_to)

        # Section 1 — OEE
        oee_levels = [self._assess_oee(l.avg_oee)[2] for l in oee.lines]
        avg_oee = sum(float(l.avg_oee) for l in oee.lines) / max(len(oee.lines), 1)
        warn_count = sum(1 for lv in oee_levels if lv == WARNING)
        crit_count = sum(1 for lv in oee_levels if lv == CRITICAL)
        oee_bullets = [
            f"• {(l.production_line or 'Все линии')}: OEE {l.avg_oee:.1f}% — {self._assess_oee(l.avg_oee)[0]}"
            for l in oee.lines
        ]
        worst_oee = self._worst(oee_levels)
        if worst_oee == NORMAL:
            oee_conclusion = (
                f"По итогам периода средний показатель OEE по предприятию составил {avg_oee:.1f}%. "
                "Все производственные линии работают в пределах целевого уровня (≥75%). "
                "Производственная эффективность соответствует установленным стандартам."
            )
            oee_rec = (
                "Рекомендуется поддерживать текущий уровень технического обслуживания и "
                "следить за стабильностью производственных процессов."
            )
        elif worst_oee == WARNING:
            oee_conclusion = (
                f"По итогам периода средний показатель OEE составил {avg_oee:.1f}%. "
                f"{warn_count} линий работают ниже целевого уровня 75%. "
                "Выявлены потери эффективности, требующие анализа и корректирующих мероприятий."
            )
            oee_rec = (
                "Рекомендуется провести анализ потерь по классификации OEE (доступность, производительность, качество) "
                "для каждой отстающей линии и разработать план улучшений."
            )
        else:
            oee_conclusion = (
                f"По итогам периода средний показатель OEE составил {avg_oee:.1f}%. "
                f"{crit_count} линий находятся в критическом состоянии (OEE < 65%). "
                "Производственная эффективность существенно ниже допустимого уровня."
            )
            oee_rec = (
                "Требуется немедленное вмешательство: проведение детального аудита потерь, "
                "задействование служб технического обслуживания, "
                "пересмотр загрузки и графиков плановых ремонтов."
            )
        self._word_section(doc, "1. ОБЩАЯ ЭФФЕКТИВНОСТЬ ОБОРУДОВАНИЯ (OEE)", oee_bullets, oee_conclusion, oee_rec,
                           chart_bytes=self._chart_gm_oee_bars(oee))

        # Section 2 — Plan
        plan_levels = [self._assess_fulfillment(l.fulfillment_pct)[2] for l in plan.lines]
        worst_plan = self._worst(plan_levels)
        plan_bullets = [
            f"• {(l.production_line or 'Без линии')}: план {l.target_quantity:.0f} / факт {l.actual_quantity:.0f} "
            f"({l.fulfillment_pct:.1f}%) — {self._assess_fulfillment(l.fulfillment_pct)[0]}"
            for l in plan.lines
        ]
        if worst_plan == NORMAL:
            plan_conclusion = (
                f"Выполнение производственного плана составило {plan.overall_fulfillment_pct:.1f}%. "
                "Все линии выполняют план на уровне ≥85%. Производство работает в штатном режиме."
            )
            plan_rec = "Рекомендуется сохранять текущую организацию производственного планирования."
        elif worst_plan == WARNING:
            plan_conclusion = (
                f"Общее выполнение плана составило {plan.overall_fulfillment_pct:.1f}%. "
                "Ряд линий имеет недовыполнение плана в диапазоне 70–85%. "
                "Выявлены риски несоблюдения производственных обязательств."
            )
            plan_rec = (
                "Рекомендуется пересмотреть оперативные планы, проверить обеспеченность "
                "сырьём и загрузку персонала на отстающих линиях."
            )
        else:
            plan_conclusion = (
                f"Общее выполнение плана составило {plan.overall_fulfillment_pct:.1f}%. "
                "Зафиксировано критическое недовыполнение плана (<70%). "
                "Производственные обязательства находятся под серьёзной угрозой."
            )
            plan_rec = (
                "Требуется срочный пересмотр производственной программы, выяснение причин срывов "
                "и оперативное совещание с руководителями линий."
            )
        self._word_section(doc, "2. ВЫПОЛНЕНИЕ ПРОИЗВОДСТВЕННОГО ПЛАНА", plan_bullets, plan_conclusion, plan_rec,
                           chart_bytes=self._chart_gm_plan_combo(plan))

        # Section 3 — Downtime
        delay_levels = [self._assess_delay(l.total_delay_hours)[2] for l in downtime.lines]
        worst_delay = self._worst(delay_levels)
        delay_bullets = [
            f"• {(l.production_line or 'Без линии')}: {l.total_delay_hours:.1f} ч простоев, "
            f"{l.delayed_orders} задержанных заказов — {self._assess_delay(l.total_delay_hours)[0]}"
            for l in downtime.lines
        ]
        if worst_delay == NORMAL:
            delay_conclusion = (
                "За анализируемый период производственные задержки отсутствуют. "
                "Все заказы завершаются в установленные сроки."
            )
            delay_rec = "Рекомендуется поддерживать текущий уровень операционной дисциплины."
        elif worst_delay == WARNING:
            delay_conclusion = (
                f"Зафиксированы незначительные задержки производства, "
                f"суммарно {downtime.total_delay_hours:.1f} ч. "
                "Ситуация не критична, однако требует внимания."
            )
            delay_rec = (
                "Рекомендуется проанализировать причины задержек и принять меры "
                "для предотвращения их повторения."
            )
        else:
            delay_conclusion = (
                f"Критический уровень простоев: суммарно {downtime.total_delay_hours:.1f} ч задержек, "
                f"{downtime.total_delayed_orders} заказов выполнено с нарушением сроков. "
                "Ситуация требует немедленного реагирования."
            )
            delay_rec = (
                "Требуется срочное расследование причин простоев, "
                "оценка потерь и разработка плана экстренных корректирующих мер."
            )
        self._word_section(doc, "3. ПРОСТОИ И ЗАДЕРЖКИ ПРОИЗВОДСТВА", delay_bullets, delay_conclusion, delay_rec,
                           chart_bytes=self._chart_gm_downtime_pareto(downtime))

        # Overall conclusion
        overall = self._worst([self._worst(oee_levels), self._worst(plan_levels), self._worst(delay_levels)])
        doc.add_heading("ИТОГОВЫЙ ВЫВОД И РЕКОМЕНДАЦИИ", 1)
        if overall == NORMAL:
            doc.add_paragraph(
                "По результатам анализа предприятие работает в штатном режиме. "
                "Все ключевые показатели находятся в норме. "
                "Рекомендуется продолжать плановую работу и мониторинг показателей."
            )
        elif overall == WARNING:
            doc.add_paragraph(
                "По результатам анализа выявлен ряд отклонений от целевых показателей. "
                "Ситуация управляема, однако требует внимания и корректирующих действий. "
                "Рекомендуется провести оперативное совещание с руководителями линий "
                "и разработать план улучшений на следующий период."
            )
        else:
            doc.add_paragraph(
                "По результатам анализа зафиксированы критические отклонения по одному или "
                "нескольким ключевым показателям. "
                "Требуется немедленное вмешательство руководства: "
                "расследование причин, оперативная корректировка планов и ресурсов, "
                "контроль исполнения мероприятий."
            )

        return self._word_to_bytes(doc)

    # ── Word: Finance ────────────────────────────────────────────────────────

    def _word_finance(
        self,
        breakdown: SalesBreakdownResponse,
        trend: RevenueTrendResponse,
        products: TopProductsResponse,
        date_from: date,
        date_to: date,
    ) -> bytes:
        from docx import Document

        doc = Document()
        self._word_header(doc, "Финансовый менеджер (Finance Manager)", date_from, date_to)

        # Section 1 — Sales breakdown
        share_levels = [self._assess_share(g.amount_share_pct)[2] for g in breakdown.groups]
        worst_share = self._worst(share_levels)
        breakdown_bullets = [
            f"• {g.group_key}: выручка {g.total_amount:,.2f}, доля {g.amount_share_pct:.1f}% — "
            f"{self._assess_share(g.amount_share_pct)[0]}"
            for g in breakdown.groups
        ]
        if worst_share == NORMAL:
            share_conclusion = (
                f"Общая выручка за период составила {breakdown.total_amount:,.2f}. "
                "Продажи распределены равномерно по каналам. Зависимость от отдельного сегмента отсутствует."
            )
            share_rec = "Рекомендуется поддерживать диверсификацию каналов сбыта."
        elif worst_share == WARNING:
            share_conclusion = (
                f"Общая выручка за период составила {breakdown.total_amount:,.2f}. "
                "Выявлена умеренная концентрация продаж в отдельных каналах (30–50%). "
                "Риски зависимости от отдельных сегментов умеренные."
            )
            share_rec = (
                "Рекомендуется следить за динамикой концентрации и постепенно "
                "диверсифицировать каналы продаж."
            )
        else:
            share_conclusion = (
                f"Общая выручка за период составила {breakdown.total_amount:,.2f}. "
                "Зафиксирована высокая концентрация продаж (>50%) в одном канале. "
                "Существенный риск зависимости."
            )
            share_rec = (
                "Рекомендуется срочно проработать стратегию диверсификации "
                "и снижения зависимости от доминирующего канала."
            )
        self._word_section(doc, "1. РАЗБИВКА ПРОДАЖ ПО КАНАЛАМ", breakdown_bullets, share_conclusion, share_rec,
                           chart_bytes=self._chart_finance_sales_pie(breakdown))

        # Section 2 — Revenue trend
        trend_levels = [self._assess_mom_growth(p.mom_growth_pct)[2] for p in trend.data]
        worst_trend = self._worst(trend_levels)
        trend_bullets = [
            f"• {p.trend_date.strftime('%d.%m.%Y')}: выручка {p.total_amount:,.2f}"
            + (f", рост MoM {p.mom_growth_pct:.1f}%" if p.mom_growth_pct is not None else "")
            for p in trend.data
        ]
        if worst_trend == NORMAL:
            trend_conclusion = "Выручка демонстрирует положительную или стабильную динамику за весь период."
            trend_rec = "Рекомендуется продолжать текущую коммерческую стратегию."
        elif worst_trend == WARNING:
            trend_conclusion = "В отдельные периоды зафиксировано снижение выручки. Динамика нестабильна."
            trend_rec = (
                "Рекомендуется проанализировать причины снижения в периоды падения "
                "и принять меры по стабилизации продаж."
            )
        else:
            trend_conclusion = "Зафиксировано критическое снижение выручки (более 10% к предыдущему периоду)."
            trend_rec = (
                "Требуется срочный анализ факторов снижения выручки "
                "и разработка антикризисного плана продаж."
            )
        self._word_section(doc, "2. ТРЕНД ВЫРУЧКИ", trend_bullets, trend_conclusion, trend_rec,
                           chart_bytes=self._chart_finance_trend_line(trend))

        # Section 3 — Top products
        prod_bullets = [
            f"• #{p.rank} {p.product_name}: выручка {p.total_amount:,.2f}, доля {p.amount_share_pct:.1f}%"
            for p in products.products
        ]
        doc.add_heading("3. ТОП ПРОДУКТОВ ПО ВЫРУЧКЕ", 2)
        self._add_chart_to_word(doc, self._chart_finance_top_barh(products))
        for b in prod_bullets:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph()

        # Overall
        overall = self._worst([worst_share, worst_trend])
        doc.add_heading("ИТОГОВЫЙ ВЫВОД И РЕКОМЕНДАЦИИ", 1)
        if overall == NORMAL:
            doc.add_paragraph(
                "Финансовые показатели предприятия за период находятся в норме. "
                "Выручка стабильна, продажи диверсифицированы. "
                "Рекомендуется поддерживать текущую коммерческую стратегию."
            )
        elif overall == WARNING:
            doc.add_paragraph(
                "Выявлены умеренные отклонения в финансовых показателях. "
                "Рекомендуется провести углублённый анализ динамики продаж "
                "и скорректировать коммерческую стратегию."
            )
        else:
            doc.add_paragraph(
                "Зафиксированы критические отклонения финансовых показателей. "
                "Требуется немедленное вмешательство: анализ причин падения, "
                "пересмотр ценовой политики и каналов сбыта."
            )

        return self._word_to_bytes(doc)

    # ── Word: QE ─────────────────────────────────────────────────────────────

    def _word_qe(
        self,
        trends: ParameterTrendsResponse,
        batches: BatchAnalysisResponse,
        pareto: DefectParetoResponse,
        date_from: date,
        date_to: date,
    ) -> bytes:
        from docx import Document

        doc = Document()
        self._word_header(doc, "Инженер по качеству (Quality Engineer)", date_from, date_to)

        # Section 1 — Parameter trends
        trend_levels = [self._assess_defect(p.overall_out_of_spec_pct)[2] for p in trends.parameters]
        worst_trend = self._worst(trend_levels)
        trend_bullets = [
            f"• {p.parameter_name}: {p.overall_out_of_spec_pct:.1f}% вне допуска "
            f"({p.total_out_of_spec} из {p.total_tests}) — {self._assess_defect(p.overall_out_of_spec_pct)[0]}"
            for p in trends.parameters
        ]
        crit_params = [p.parameter_name for p in trends.parameters if self._assess_defect(p.overall_out_of_spec_pct)[2] == CRITICAL]
        if worst_trend == NORMAL:
            trend_conclusion = (
                "Все контролируемые параметры качества находятся в пределах допустимых значений. "
                "Уровень несоответствий не превышает 5%."
            )
            trend_rec = "Рекомендуется поддерживать текущий уровень контроля качества."
        elif worst_trend == WARNING:
            trend_conclusion = (
                "По ряду параметров качества зафиксированы повышенные уровни несоответствий (5–10%). "
                "Ситуация требует внимания и анализа первопричин."
            )
            trend_rec = (
                "Рекомендуется провести анализ первопричин по параметрам с повышенным уровнем "
                "несоответствий и скорректировать технологические параметры."
            )
        else:
            params_str = ", ".join(crit_params) if crit_params else "ряде параметров"
            trend_conclusion = (
                f"Критический уровень несоответствий (>10%) зафиксирован по: {params_str}. "
                "Качество продукции не соответствует установленным требованиям."
            )
            trend_rec = (
                "Требуется немедленная остановка или усиленный контроль по критическим параметрам, "
                "инициирование корректирующих действий и уведомление руководства."
            )
        self._word_section(doc, "1. ТРЕНДЫ ПАРАМЕТРОВ КАЧЕСТВА", trend_bullets, trend_conclusion, trend_rec,
                           chart_bytes=self._chart_qe_params_bars(trends))

        # Section 2 — Batch analysis
        batch_bullets = [
            f"• {l.lot_number} ({l.product_name or 'н/д'}): {l.fail_count} нарушений"
            + (f", дата {l.production_date.strftime('%d.%m.%Y')}" if l.production_date else "")
            for l in batches.lots
        ]
        batch_conclusion: str
        batch_rec: str
        if batches.lot_count == 0:
            batch_conclusion = "За анализируемый период партий с отклонениями не выявлено."
            batch_rec = "Рекомендуется сохранять текущий уровень входного и операционного контроля."
        elif batches.lot_count <= 3:
            batch_conclusion = (
                f"Выявлено {batches.lot_count} партий с отклонениями от норм качества. "
                "Количество незначительное."
            )
            batch_rec = "Рекомендуется провести анализ причин отклонений по каждой партии."
        else:
            batch_conclusion = (
                f"Выявлено {batches.lot_count} партий с отклонениями — повышенное количество. "
                "Системная проблема качества требует расследования."
            )
            batch_rec = (
                "Требуется корневой анализ причин (RCA), оценка соответствия сырья "
                "и пересмотр технологических параметров."
            )
        self._word_section(doc, "2. АНАЛИЗ ПАРТИЙ С ОТКЛОНЕНИЯМИ", batch_bullets, batch_conclusion, batch_rec,
                           chart_bytes=self._chart_qe_batches_timeline(batches))

        # Section 3 — Pareto
        if pareto.parameters:
            top3 = pareto.parameters[:3]
            pareto_bullets = [
                f"• {p.parameter_name}: {p.defect_count} дефектов ({p.defect_pct:.1f}%), кумул. {p.cumulative_pct:.1f}%"
                for p in top3
            ]
        else:
            pareto_bullets = ["• Дефекты за период не зафиксированы"]

        doc.add_heading("3. PARETO-АНАЛИЗ ДЕФЕКТОВ", 2)
        self._add_chart_to_word(doc, self._chart_qe_pareto(pareto))
        for b in pareto_bullets:
            doc.add_paragraph(b, style="List Bullet")
        p = doc.add_paragraph()
        p.add_run("Вывод: ").bold = True
        p.add_run(
            f"Всего за период зафиксировано {pareto.total_defects} дефектов. "
            "Наибольший вклад вносят параметры, указанные выше."
        )
        doc.add_paragraph()

        # Overall
        doc.add_heading("ИТОГОВЫЙ ВЫВОД И РЕКОМЕНДАЦИИ", 1)
        overall = self._worst(trend_levels)
        if overall == NORMAL:
            doc.add_paragraph(
                "Система качества функционирует в штатном режиме. "
                "Все ключевые параметры в норме. Дефектность на приемлемом уровне."
            )
        elif overall == WARNING:
            doc.add_paragraph(
                "Выявлены умеренные отклонения в отдельных параметрах качества. "
                "Рекомендуется анализ первопричин и корректирующие мероприятия."
            )
        else:
            doc.add_paragraph(
                "Критические отклонения в системе качества требуют немедленного реагирования. "
                "Необходимо задействовать процедуры управления несоответствующей продукцией "
                "и провести внеплановый аудит системы качества."
            )

        return self._word_to_bytes(doc)

    # ── Word: Line Master ────────────────────────────────────────────────────

    def _word_line_master(
        self,
        progress: ShiftProgressResponse,
        comparison: ShiftComparisonResponse,
        defects: DefectSummaryResponse,
        production_date: date,
        date_from: date,
        date_to: date,
    ) -> bytes:
        from docx import Document

        doc = Document()
        self._word_header(doc, "Мастер линии (Line Master)", date_from, date_to)

        # Section 1 — Shift progress
        shift_levels = [self._assess_defect(s.defect_rate)[2] for s in progress.shifts]
        worst_shift = self._worst(shift_levels)
        shift_bullets = [
            f"• {s.shift}: {s.lot_count} лотов, {s.total_quantity:.1f} кг, "
            f"дефектность {s.defect_rate:.1f}% — {self._assess_defect(s.defect_rate)[0]}"
            for s in progress.shifts
        ]
        if worst_shift == NORMAL:
            shift_conclusion = (
                f"На {production_date.strftime('%d.%m.%Y')} все смены работают в штатном режиме. "
                "Дефектность по всем сменам не превышает 5%."
            )
            shift_rec = "Рекомендуется поддерживать текущий уровень операционного контроля."
        elif worst_shift == WARNING:
            shift_conclusion = (
                f"На {production_date.strftime('%d.%m.%Y')} зафиксирован повышенный уровень дефектов "
                "в отдельных сменах (5–10%). Требует внимания."
            )
            shift_rec = (
                "Рекомендуется провести инструктаж операторов и проверить "
                "технологические параметры линии."
            )
        else:
            shift_conclusion = (
                f"На {production_date.strftime('%d.%m.%Y')} зафиксирован критический уровень дефектов "
                "в одной или нескольких сменах (>10%)."
            )
            shift_rec = (
                "Требуется немедленная проверка оборудования, сырья и соблюдения "
                "технологического регламента. Уведомить ОТК."
            )
        self._word_section(doc, "1. ПРОГРЕСС СМЕН", shift_bullets, shift_conclusion, shift_rec,
                           chart_bytes=self._chart_lm_shift_combo(progress))

        # Section 2 — Shift comparison (added)
        doc.add_heading("2. СРАВНЕНИЕ СМЕН ЗА ПЕРИОД", 2)
        self._add_chart_to_word(doc, self._chart_lm_comparison_line(comparison))
        comp_bullets = [
            f"• {p.date.strftime('%d.%m.%Y')} ({p.shift or 'День'}): "
            f"{p.lot_count} лотов, {p.total_quantity:.1f} кг, {p.defect_count} дефектов"
            for p in comparison.shifts
        ]
        for b in comp_bullets:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph()

        # Section 3 — Defect summary
        defect_levels = [self._assess_defect(i.fail_rate)[2] for i in defects.items]
        worst_defect = self._worst(defect_levels)
        defect_bullets = [
            f"• {i.parameter_name}: {i.failed_tests} из {i.total_tests} тестов провалено "
            f"({i.fail_rate:.1f}%) — {self._assess_defect(i.fail_rate)[0]}"
            for i in defects.items
        ]
        if worst_defect == NORMAL:
            defect_conclusion = (
                "За анализируемый период все параметры качества находятся в допустимых пределах."
            )
            defect_rec = "Рекомендуется продолжать плановый контроль качества."
        elif worst_defect == WARNING:
            defect_conclusion = (
                f"По ряду параметров зафиксирован повышенный уровень отказов. "
                f"Всего за период: {defects.total_defects} случаев несоответствия."
            )
            defect_rec = "Рекомендуется ужесточить контроль по параметрам с повышенным уровнем отказов."
        else:
            defect_conclusion = (
                f"Критический уровень несоответствий: {defects.total_defects} случаев за период. "
                "Качество продукции под угрозой."
            )
            defect_rec = (
                "Требуется немедленная эскалация ситуации, остановка отгрузки "
                "и детальный разбор причин несоответствий."
            )
        self._word_section(doc, "3. СВОДКА ДЕФЕКТОВ ПО ПАРАМЕТРАМ", defect_bullets, defect_conclusion, defect_rec,
                           chart_bytes=self._chart_lm_defects_barh(defects))

        # Overall
        overall = self._worst([worst_shift, worst_defect])
        doc.add_heading("ИТОГОВЫЙ ВЫВОД И РЕКОМЕНДАЦИИ", 1)
        if overall == NORMAL:
            doc.add_paragraph(
                "Производство работает в штатном режиме. "
                "Качественные показатели соответствуют нормам. "
                "Рекомендуется продолжать плановую работу."
            )
        elif overall == WARNING:
            doc.add_paragraph(
                "Выявлены умеренные отклонения. Ситуация управляема, "
                "однако требует усиленного контроля и корректирующих действий."
            )
        else:
            doc.add_paragraph(
                "Зафиксированы критические отклонения. "
                "Необходимо незамедлительно сообщить руководству, "
                "выявить и устранить причины нарушений."
            )

        return self._word_to_bytes(doc)

    # ── Production Overview Export ───────────────────────────────────────────

    async def export_production_overview(
        self,
        fmt: str,
        date_from: date,
        date_to: date,
        production_line_id: Optional[str] = None,
    ) -> bytes:
        from app.services.production_analytics_service import ProductionAnalyticsService
        from app.services.order_service import OrderService
        from app.services.output_service import OutputService
        from app.services.quality_service import QualityService
        from app.services.sales_service import SalesService
        from app.services.sensor_service import SensorService
        from app.services.inventory_service import InventoryService

        pa_svc = ProductionAnalyticsService(self.db)
        kpi = await pa_svc.get_kpi(date_from, date_to, production_line_id, compare_with_previous=True)
        otif = await pa_svc.get_otif(date_from, date_to, production_line_id)

        order_svc = OrderService(self.db)
        orders = await order_svc.get_status_summary(date_from, date_to, production_line_id)

        output_svc = OutputService(self.db)
        shifts = await output_svc.get_output_by_shift(date_from, date_to)

        q_svc = QualityService(self.db)
        q_summary = await q_svc.get_quality_summary(date_from, date_to)
        q_lots = await q_svc.get_quality_lots(date_from, date_to)

        sales_svc = SalesService(self.db)
        regions = await sales_svc.get_sales_by_regions(date_from, date_to)

        sensor_svc = SensorService(self.db)
        sensors = await sensor_svc.get_sensor_stats(production_line_id)

        inv_svc = InventoryService(self.db)
        inventory = await inv_svc.get_current_inventory()

        logger.info("export_production_overview", fmt=fmt, date_from=date_from, date_to=date_to)
        if fmt == "docx":
            return self._word_production_overview(
                kpi, otif, orders, shifts, q_summary, q_lots, regions, sensors, inventory, date_from, date_to
            )
        return self._excel_production_overview(
            kpi, otif, orders, shifts, q_summary, q_lots, regions, sensors, inventory, date_from, date_to
        )

    # ── Charts: Production Overview ──────────────────────────────────────────

    def _chart_po_kpi_bars(self, kpi: dict, otif: dict) -> bytes:
        total_orders = int(kpi.get("total_orders", 0))
        completed_orders = int(kpi.get("completed_orders", 0))
        fulfillment_pct = (completed_orders / total_orders * 100) if total_orders else 0.0
        metrics = [
            ("OEE", float(kpi.get("oee_estimate", 0)) * 100, 85.0),
            ("Выполнение плана", fulfillment_pct, 90.0),
            ("OTIF", float(otif.get("otif_rate", 0)) * 100, 95.0),
        ]
        labels = [m[0] for m in metrics]
        values = [m[1] for m in metrics]
        targets = [m[2] for m in metrics]

        fig, ax = plt.subplots(figsize=(10, 5))
        y = list(range(len(labels)))
        bars = ax.barh(y, values, color="#366092", edgecolor="white", linewidth=0.5, zorder=2, label="Факт")
        ax.scatter(targets, y, color="#FF4444", zorder=4, s=80, marker="|", linewidths=3, label="Цель")
        ax.set_yticks(y)
        ax.set_yticklabels(labels, fontsize=11)
        ax.set_xlabel("%")
        ax.set_xlim(0, 115)
        ax.set_title("Ключевые производственные показатели", fontsize=13, fontweight="bold", pad=15)
        ax.legend(fontsize=10)
        ax.grid(axis="x", alpha=0.3, zorder=1)
        for bar, val in zip(bars, values):
            ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
                    f"{val:.1f}%", va="center", fontsize=10)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=100, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_po_orders_stacked(self, orders: Any) -> bytes:
        by_line = getattr(orders, "by_production_line", {}) or {}
        if not by_line:
            return self._no_data_chart("Статус заказов по линиям")

        lines = list(by_line.keys())
        statuses = ["planned", "in_progress", "completed", "cancelled"]
        labels_ru = {"planned": "Плановые", "in_progress": "В работе",
                     "completed": "Завершённые", "cancelled": "Отменённые"}
        colors_map = {"planned": "#4472C4", "in_progress": "#FFC000",
                      "completed": "#92D050", "cancelled": "#FF4444"}
        x = list(range(len(lines)))
        bottoms = [0.0] * len(lines)

        fig, ax = plt.subplots(figsize=(10, 6))
        for status in statuses:
            vals = [float((by_line[ln] or {}).get(status, 0) if isinstance(by_line[ln], dict)
                          else getattr(by_line[ln], status, 0)) for ln in lines]
            ax.bar(x, vals, bottom=bottoms, label=labels_ru[status],
                   color=colors_map[status], edgecolor="white", linewidth=0.5, zorder=2)
            bottoms = [bottoms[i] + vals[i] for i in range(len(lines))]

        ax.set_xticks(x)
        ax.set_xticklabels(lines, rotation=30, ha="right", fontsize=9)
        ax.set_ylabel("Заказов")
        ax.set_title("Статус заказов по производственным линиям", fontsize=13, fontweight="bold", pad=15)
        ax.legend(fontsize=9)
        ax.grid(axis="y", alpha=0.3, zorder=1)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=100, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_po_shift_combo(self, shifts: dict) -> bytes:
        items = shifts.get("items", [])
        if not items:
            return self._no_data_chart("Выпуск продукции по сменам")

        from collections import defaultdict
        by_date: dict = defaultdict(float)
        by_shift: dict = defaultdict(lambda: defaultdict(float))
        for item in items:
            d = str(item["date"])
            s = item["shift"]
            q = float(item.get("total_quantity", 0))
            by_date[d] += q
            by_shift[s][d] += q

        dates = sorted(by_date.keys())
        shift_names = sorted(by_shift.keys())
        x = list(range(len(dates)))
        shift_colors = ["#4472C4", "#FFC000", "#92D050", "#FF4444"]

        fig, ax1 = plt.subplots(figsize=(10, 6))
        ax2 = ax1.twinx()
        w = 0.6 / max(len(shift_names), 1)
        for i, shift in enumerate(shift_names):
            vals = [by_shift[shift].get(d, 0.0) for d in dates]
            offset = (i - len(shift_names) / 2 + 0.5) * w
            ax1.bar([xi + offset for xi in x], vals, width=w,
                    label=shift, color=shift_colors[i % len(shift_colors)],
                    edgecolor="white", linewidth=0.3, zorder=2, alpha=0.85)

        totals = [by_date[d] for d in dates]
        ax2.plot(x, totals, color="#FF4444", linewidth=2, marker="o", markersize=5,
                 label="Итого", zorder=3)
        ax1.set_xticks(x)
        ax1.set_xticklabels(dates, rotation=45, ha="right", fontsize=8)
        ax1.set_ylabel("Объём (ед.)")
        ax2.set_ylabel("Итого по дню", color="#FF4444")
        ax1.set_title("Выпуск продукции по сменам", fontsize=13, fontweight="bold", pad=15)
        lines1, labs1 = ax1.get_legend_handles_labels()
        lines2, labs2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, labs1 + labs2, fontsize=9)
        ax1.grid(axis="y", alpha=0.3, zorder=1)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=100, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_po_quality_params(self, q_summary: Any) -> bytes:
        by_param = getattr(q_summary, "by_parameter", {}) or {}
        if not by_param:
            return self._no_data_chart("Качество по параметрам")

        params = list(by_param.keys())
        values = [float(by_param[p].get("in_spec_percent", 0) if isinstance(by_param[p], dict)
                        else getattr(by_param[p], "in_spec_percent", 0)) for p in params]
        colors = [f"#{STATUS_COLOR[self._assess_defect(Decimal(str((100 - v) / 100)))[2]]}" for v in values]

        fig, ax = plt.subplots(figsize=(10, 5))
        bars = ax.bar(params, values, color=colors, edgecolor="white", linewidth=0.5, zorder=2)
        ax.axhline(y=95, color="#366092", linestyle="--", linewidth=2, label="Цель 95%", zorder=3)
        ax.set_ylabel("% в норме")
        ax.set_ylim(0, 110)
        ax.set_title("Параметры качества — доля в норме", fontsize=13, fontweight="bold", pad=15)
        ax.legend(fontsize=10)
        ax.grid(axis="y", alpha=0.3, zorder=1)
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1,
                    f"{val:.1f}%", ha="center", va="bottom", fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=100, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    def _chart_po_sales_pie(self, regions: Any) -> bytes:
        region_list = getattr(regions, "regions", []) or []
        if not region_list:
            return self._no_data_chart("Продажи по регионам")

        labels = [r.get("region", "—") if isinstance(r, dict) else getattr(r, "region", "—")
                  for r in region_list]
        values = [float(r.get("total_amount", 0) if isinstance(r, dict) else getattr(r, "total_amount", 0))
                  for r in region_list]

        fig, ax = plt.subplots(figsize=(9, 6))
        _, _, autotexts_po = ax.pie(  # type: ignore[misc]
            values, labels=labels, autopct="%1.1f%%",
            startangle=90, pctdistance=0.8,
            wedgeprops={"edgecolor": "white", "linewidth": 1},
        )
        for at in autotexts_po:
            at.set_fontsize(9)
        ax.set_title("Продажи по регионам", fontsize=13, fontweight="bold", pad=15)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=100, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    # ── Excel: Production Overview ────────────────────────────────────────────

    def _excel_production_overview(
        self,
        kpi: dict,
        otif: dict,
        orders: Any,
        shifts: dict,
        q_summary: Any,
        q_lots: Any,
        regions: Any,
        sensors: dict,
        inventory: dict,
        date_from: date,
        date_to: date,
    ) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()

        # ── Sheet 1: KPI и OTIF ──────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "KPI и OTIF"
        headers1 = ["Метрика", "Значение", "Цель", "Статус", "Изм. к пред. периоду"]
        self._style_header(ws1, headers1)

        total_orders = int(kpi.get("total_orders", 0))
        completed_orders = int(kpi.get("completed_orders", 0))
        fulfillment_pct = (completed_orders / total_orders * 100) if total_orders else 0.0
        oee_pct = float(kpi.get("oee_estimate", 0)) * 100
        defect_pct = float(kpi.get("defect_rate", 0)) * 100
        change = kpi.get("change_percent") or {}

        kpi_rows: list[tuple] = [
            ("Общий выпуск (ед.)", float(kpi.get("total_output", 0)), "—",
             self._assess_fulfillment(Decimal(str(fulfillment_pct)))[0],
             f"{change.get('total_output', 0):+.1f}%" if change else "—"),
            ("OEE (%)", round(oee_pct, 1), 85.0,
             self._assess_oee(Decimal(str(oee_pct)))[0],
             f"{change.get('oee_estimate', 0):+.1f}%" if change else "—"),
            ("Доля дефектов (%)", round(defect_pct, 2), "≤ 1.5",
             self._assess_defect(kpi.get("defect_rate", Decimal("0")))[0], "—"),
            ("Выполнение плана (%)", round(fulfillment_pct, 1), 90.0,
             self._assess_fulfillment(Decimal(str(fulfillment_pct)))[0], "—"),
        ]
        for row_vals in kpi_rows:
            label = row_vals[0]
            if "OEE" in label:
                _, color, _ = self._assess_oee(Decimal(str(row_vals[1])))
            elif "дефект" in label.lower():
                _, color, _ = self._assess_defect(kpi.get("defect_rate", Decimal("0")))
            else:
                _, color, _ = self._assess_fulfillment(Decimal(str(
                    row_vals[1] if isinstance(row_vals[1], (int, float)) else 0
                )))
            ws1.append(list(row_vals))
            self._color_row(ws1, ws1.max_row, color, len(headers1))

        otif_pct = float(otif.get("otif_rate", 0)) * 100
        otif_total = int(otif.get("total_orders", 0))
        on_time_pct = (float(otif.get("on_time_orders", 0)) / otif_total * 100) if otif_total else 0.0
        otif_rows: list[tuple] = [
            ("OTIF Rate (%)", round(otif_pct, 1), 95.0,
             self._assess_fulfillment(Decimal(str(otif_pct)))[0], "—"),
            ("В срок (%)", round(on_time_pct, 1), 95.0,
             self._assess_fulfillment(Decimal(str(on_time_pct)))[0], "—"),
            ("В полном объёме (заказов)", otif.get("in_full_quantity_orders", 0), "—", "—", "—"),
            ("OTIF-заказов / всего",
             f"{otif.get('otif_orders', 0)} / {otif.get('total_orders', 0)}", "—", "—", "—"),
        ]
        for row_vals in otif_rows:
            ws1.append(list(row_vals))

        self._add_title(ws1, f"KPI и OTIF | {date_from} — {date_to}", len(headers1))
        self._autofit(ws1)
        self._add_chart_to_excel(ws1, self._chart_po_kpi_bars(kpi, otif), "G2")

        # ── Sheet 2: Заказы по линиям ────────────────────────────────────────
        ws2 = wb.create_sheet("Заказы по линиям")
        headers2 = ["Линия", "Плановые", "В работе", "Завершённые", "Отменённые", "Итого", "Выполнение %"]
        self._style_header(ws2, headers2)

        by_line = getattr(orders, "by_production_line", {}) or {}
        for line_name, counts in by_line.items():
            planned = counts.get("planned", 0) if isinstance(counts, dict) else getattr(counts, "planned", 0)
            in_prog = counts.get("in_progress", 0) if isinstance(counts, dict) else getattr(counts, "in_progress", 0)
            completed = counts.get("completed", 0) if isinstance(counts, dict) else getattr(counts, "completed", 0)
            cancelled = counts.get("cancelled", 0) if isinstance(counts, dict) else getattr(counts, "cancelled", 0)
            total = planned + in_prog + completed + cancelled
            exec_pct = round(completed / total * 100, 1) if total else 0.0
            _, color, _ = self._assess_fulfillment(Decimal(str(exec_pct)))
            ws2.append([line_name, planned, in_prog, completed, cancelled, total, exec_pct])
            self._color_row(ws2, ws2.max_row, color, len(headers2))

        by_status = getattr(orders, "by_status", None)
        planned_t = getattr(by_status, "planned", 0) if by_status else 0
        in_prog_t = getattr(by_status, "in_progress", 0) if by_status else 0
        completed_t = getattr(by_status, "completed", 0) if by_status else 0
        cancelled_t = getattr(by_status, "cancelled", 0) if by_status else 0
        total_all = planned_t + in_prog_t + completed_t + cancelled_t
        self._add_summary(ws2, [
            "ИТОГО",
            planned_t,
            in_prog_t,
            completed_t,
            cancelled_t,
            total_all,
            "",
        ])
        self._add_title(ws2, f"Статус заказов по линиям | {date_from} — {date_to}", len(headers2))
        self._autofit(ws2)
        self._add_chart_to_excel(ws2, self._chart_po_orders_stacked(orders), "I2")

        # ── Sheet 3: Выпуск по сменам ────────────────────────────────────────
        ws3 = wb.create_sheet("Выпуск по сменам")
        headers3 = ["Дата", "Смена", "Объём (ед.)"]
        self._style_header(ws3, headers3)

        shift_items = shifts.get("items", [])
        total_qty = 0.0
        for item in shift_items:
            qty = float(item.get("total_quantity", 0))
            total_qty += qty
            ws3.append([str(item["date"]), item.get("shift", "—"), round(qty, 1)])

        self._add_summary(ws3, ["ИТОГО", "", round(total_qty, 1)])
        self._add_title(ws3, f"Выпуск по сменам | {date_from} — {date_to}", len(headers3))
        self._autofit(ws3)
        self._add_chart_to_excel(ws3, self._chart_po_shift_combo(shifts), "E2")

        # ── Sheet 4: Качество ────────────────────────────────────────────────
        ws4 = wb.create_sheet("Качество")

        # Summary block
        headers4a = ["Метрика", "Значение"]
        self._style_header(ws4, headers4a)
        avg_q = float(getattr(q_summary, "average_quality", 0))
        def_r = float(getattr(q_summary, "defect_rate", 0))
        for row_vals in [
            ("Средний балл качества (%)", round(avg_q, 1)),
            ("Доля дефектов (%)", round(def_r, 1)),
            ("Одобрено (партий)", getattr(q_summary, "approved_count", 0)),
            ("Отклонено (партий)", getattr(q_summary, "rejected_count", 0)),
        ]:
            ws4.append(list(row_vals))
        ws4.append([])

        # Parameters block
        self._style_header(ws4, ["Параметр", "% в норме", "Тестов"])
        by_param = getattr(q_summary, "by_parameter", {}) or {}
        for param_name, stats in by_param.items():
            in_spec = float(stats.get("in_spec_percent", 0) if isinstance(stats, dict)
                            else getattr(stats, "in_spec_percent", 0))
            tests = int(stats.get("tests_count", 0) if isinstance(stats, dict)
                        else getattr(stats, "tests_count", 0))
            _, color, _ = self._assess_defect(Decimal(str((100 - in_spec) / 100)))
            ws4.append([param_name, round(in_spec, 1), tests])
            self._color_row(ws4, ws4.max_row, color, 3)
        ws4.append([])

        # Lots block
        lots_hdr = ["Лот", "Продукт", "Дата", "Протестировано", "Прошло", "% прохождения", "Решение"]
        self._style_header(ws4, lots_hdr)
        decision_color = {"approved": COLOR_GREEN, "rejected": COLOR_RED, "pending": COLOR_YELLOW}
        for lot in (getattr(q_lots, "lots", []) or []):
            tested = int(lot.get("parameters_tested", 0) if isinstance(lot, dict)
                         else getattr(lot, "parameters_tested", 0))
            passed = int(lot.get("parameters_passed", 0) if isinstance(lot, dict)
                         else getattr(lot, "parameters_passed", 0))
            pct = round(passed / tested * 100, 1) if tested else 0.0
            decision = (lot.get("decision", "pending") if isinstance(lot, dict)
                        else getattr(lot, "decision", "pending")) or "pending"
            ws4.append([
                lot.get("lot_number", "") if isinstance(lot, dict) else getattr(lot, "lot_number", ""),
                lot.get("product_name", "") if isinstance(lot, dict) else getattr(lot, "product_name", ""),
                str(lot.get("test_date", "") if isinstance(lot, dict) else getattr(lot, "test_date", "")),
                tested, passed, pct, decision,
            ])
            self._color_row(ws4, ws4.max_row, decision_color.get(decision, COLOR_YELLOW), len(lots_hdr))

        self._add_title(ws4, f"Контроль качества | {date_from} — {date_to}", len(lots_hdr))
        self._autofit(ws4)
        self._add_chart_to_excel(ws4, self._chart_po_quality_params(q_summary), "I2")

        # ── Sheet 5: Продажи, Сенсоры, Запасы ───────────────────────────────
        ws5 = wb.create_sheet("Продажи, Сенсоры, Запасы")

        # Sales
        self._style_header(ws5, ["Регион", "Выручка (руб.)", "Объём", "Сделок", "Доля %"])
        region_list = getattr(regions, "regions", []) or []
        for r in region_list:
            ws5.append([
                r.get("region", "—") if isinstance(r, dict) else getattr(r, "region", "—"),
                float(r.get("total_amount", 0) if isinstance(r, dict) else getattr(r, "total_amount", 0)),
                float(r.get("total_quantity", 0) if isinstance(r, dict) else getattr(r, "total_quantity", 0)),
                int(r.get("sales_count", 0) if isinstance(r, dict) else getattr(r, "sales_count", 0)),
                float(r.get("percentage", 0) if isinstance(r, dict) else getattr(r, "percentage", 0)),
            ])
        ws5.append([])

        # Sensors
        sensor_hdr = ["Параметр", "Линия", "Среднее", "Мин", "Макс", "Ед.", "Тревог", "Статус"]
        self._style_header(ws5, sensor_hdr)
        for s in (sensors.get("items", []) or []):
            alert_count = int(s.get("alert_count", 0) or 0)
            reading_count = int(s.get("reading_count", 0) or 1)
            alert_ratio = alert_count / reading_count
            s_label, s_color, _ = self._assess_defect(Decimal(str(round(alert_ratio, 4))))
            ws5.append([
                s.get("parameter_name", "—"),
                s.get("production_line_id", "—"),
                round(float(s.get("avg_value") or 0), 2),
                round(float(s.get("min_value") or 0), 2),
                round(float(s.get("max_value") or 0), 2),
                s.get("unit", "—"),
                alert_count,
                s_label,
            ])
            self._color_row(ws5, ws5.max_row, s_color, len(sensor_hdr))
        ws5.append([])

        # Inventory
        self._style_header(ws5, ["Продукт", "Склад", "Лот", "Кол-во", "Ед. изм."])
        snap_date = inventory.get("snapshot_date", "—")
        for item in (inventory.get("items", []) or []):
            ws5.append([
                item.get("product_name", "—"),
                item.get("warehouse_code", "—"),
                item.get("lot_number", "—"),
                round(float(item.get("quantity", 0)), 2),
                item.get("unit_of_measure", "—"),
            ])

        self._add_title(ws5, f"Продажи, Сенсоры, Запасы | срез: {snap_date}", len(sensor_hdr))
        self._autofit(ws5)
        self._add_chart_to_excel(ws5, self._chart_po_sales_pie(regions), "G2")

        return self._wb_to_bytes(wb)

    # ── Word: Production Overview ─────────────────────────────────────────────

    def _word_production_overview(
        self,
        kpi: dict,
        otif: dict,
        orders: Any,
        shifts: dict,
        q_summary: Any,
        q_lots: Any,
        regions: Any,
        sensors: dict,
        inventory: dict,
        date_from: date,
        date_to: date,
    ) -> bytes:
        from docx import Document

        doc = Document()
        self._word_header(doc, "Сводный отчёт по производству", date_from, date_to)

        # ── Section 1: KPI и OTIF ────────────────────────────────────────────
        total_orders = int(kpi.get("total_orders", 0))
        completed_orders = int(kpi.get("completed_orders", 0))
        fulfillment_pct = (completed_orders / total_orders * 100) if total_orders else 0.0
        oee_pct = float(kpi.get("oee_estimate", 0)) * 100
        defect_pct = float(kpi.get("defect_rate", 0)) * 100
        otif_pct = float(otif.get("otif_rate", 0)) * 100
        change = kpi.get("change_percent") or {}

        kpi_bullets = [
            f"• Общий выпуск: {float(kpi.get('total_output', 0)):.0f} ед."
            + (f" ({change.get('total_output', 0):+.1f}% к пред. периоду)" if change else ""),
            f"• OEE: {oee_pct:.1f}% — {self._assess_oee(Decimal(str(oee_pct)))[0]}"
            + (f" ({change.get('oee_estimate', 0):+.1f}%)" if change else ""),
            f"• Доля дефектов: {defect_pct:.2f}% — {self._assess_defect(kpi.get('defect_rate', Decimal('0')))[0]}",
            f"• Выполнение плана: {fulfillment_pct:.1f}% ({completed_orders}/{total_orders} заказов)",
            f"• OTIF: {otif_pct:.1f}% (в срок: {otif.get('on_time_orders', 0)}, "
            f"в полном объёме: {otif.get('in_full_quantity_orders', 0)} из {otif.get('total_orders', 0)})",
        ]
        worst_kpi = self._worst([
            self._assess_oee(Decimal(str(oee_pct)))[2],
            self._assess_defect(kpi.get("defect_rate", Decimal("0")))[2],
            self._assess_fulfillment(Decimal(str(fulfillment_pct)))[2],
        ])
        if worst_kpi == NORMAL:
            kpi_conclusion = (
                f"Производство работает эффективно: OEE {oee_pct:.1f}%, "
                f"доля дефектов {defect_pct:.2f}%, выполнение плана {fulfillment_pct:.1f}%. "
                "Все показатели соответствуют целевым значениям."
            )
            kpi_rec = "Рекомендуется поддерживать текущий уровень операционной эффективности."
        elif worst_kpi == WARNING:
            kpi_conclusion = (
                f"Выявлены умеренные отклонения: OEE {oee_pct:.1f}%, "
                f"выполнение плана {fulfillment_pct:.1f}%. "
                "Показатели требуют внимания для предотвращения дальнейшего ухудшения."
            )
            kpi_rec = (
                "Рекомендуется провести анализ узких мест по линиям, "
                "проверить загрузку оборудования и сбалансировать смены."
            )
        else:
            kpi_conclusion = (
                f"Зафиксированы критические отклонения: OEE {oee_pct:.1f}%, "
                f"доля дефектов {defect_pct:.2f}%, выполнение плана {fulfillment_pct:.1f}%. "
                "Ситуация требует немедленного реагирования."
            )
            kpi_rec = (
                "Необходимо срочно выявить причины снижения эффективности, "
                "инициировать технический аудит и пересмотреть производственный план."
            )
        self._word_section(doc, "1. KPI И OTIF", kpi_bullets, kpi_conclusion, kpi_rec,
                           chart_bytes=self._chart_po_kpi_bars(kpi, otif))

        # ── Section 2: Статус заказов ────────────────────────────────────────
        by_line = getattr(orders, "by_production_line", {}) or {}
        order_bullets: list[str] = []
        order_levels: list[int] = []
        for line_name, counts in by_line.items():
            compl = counts.get("completed", 0) if isinstance(counts, dict) else getattr(counts, "completed", 0)
            total = sum(
                (counts.get(s, 0) if isinstance(counts, dict) else getattr(counts, s, 0))
                for s in ["planned", "in_progress", "completed", "cancelled"]
            )
            pct = (compl / total * 100) if total else 0.0
            _, _, level = self._assess_fulfillment(Decimal(str(pct)))
            order_levels.append(level)
            order_bullets.append(
                f"• {line_name}: {compl}/{total} заказов выполнено ({pct:.1f}%) — "
                f"{self._assess_fulfillment(Decimal(str(pct)))[0]}"
            )
        if not order_bullets:
            order_bullets = ["• Данные по заказам в разрезе линий отсутствуют."]
        worst_orders = self._worst(order_levels) if order_levels else NORMAL
        if worst_orders == NORMAL:
            ord_conclusion = "Все производственные линии выполняют план в соответствии с установленными нормами."
            ord_rec = "Рекомендуется сохранять текущий ритм производства и контролировать очередь заказов."
        elif worst_orders == WARNING:
            ord_conclusion = "Ряд производственных линий испытывает трудности с выполнением плановых показателей."
            ord_rec = "Рекомендуется перераспределить загрузку и усилить контроль за ходом выполнения заказов."
        else:
            ord_conclusion = "Критическое отставание по выполнению заказов на одной или нескольких линиях."
            ord_rec = (
                "Необходимо срочное вмешательство: анализ причин задержек, "
                "приоритизация критических заказов, при необходимости — сверхурочная работа."
            )
        self._word_section(doc, "2. СТАТУС ЗАКАЗОВ ПО ПРОИЗВОДСТВЕННЫМ ЛИНИЯМ",
                           order_bullets, ord_conclusion, ord_rec,
                           chart_bytes=self._chart_po_orders_stacked(orders))

        # ── Section 3: Выпуск по сменам и Качество ──────────────────────────
        shift_items = shifts.get("items", [])
        total_output_qty = sum(float(i.get("total_quantity", 0)) for i in shift_items)
        shift_names = sorted({i.get("shift", "—") for i in shift_items})
        shift_totals = {
            s: sum(float(i.get("total_quantity", 0)) for i in shift_items if i.get("shift") == s)
            for s in shift_names
        }
        shift_bullets: list[str] = [f"• Общий выпуск за период: {total_output_qty:.0f} ед."]
        for s, qty in shift_totals.items():
            pct = (qty / total_output_qty * 100) if total_output_qty else 0
            shift_bullets.append(f"• {s}: {qty:.0f} ед. ({pct:.1f}%)")

        avg_q = float(getattr(q_summary, "average_quality", 0))
        def_r = float(getattr(q_summary, "defect_rate", 0))
        approved = int(getattr(q_summary, "approved_count", 0))
        rejected = int(getattr(q_summary, "rejected_count", 0))
        shift_bullets += [
            f"• Средний балл качества: {avg_q:.1f}%",
            f"• Доля дефектов: {def_r:.1f}%",
            f"• Одобрено / отклонено партий: {approved} / {rejected}",
        ]
        worst_q = self._assess_defect(getattr(q_summary, "defect_rate", Decimal("0")))[2]
        if worst_q == NORMAL:
            q_conclusion = (
                f"Качество продукции соответствует стандартам: "
                f"средний балл {avg_q:.1f}%, доля дефектов {def_r:.1f}%."
            )
            q_rec = "Рекомендуется сохранять текущий контроль качества и регулярно проводить аудит параметров."
        elif worst_q == WARNING:
            q_conclusion = f"Доля дефектов {def_r:.1f}% превышает норму. Требуется анализ по параметрам."
            q_rec = "Рекомендуется усилить входной контроль сырья и проверить соблюдение технологических регламентов."
        else:
            q_conclusion = (
                f"Критический уровень дефектов {def_r:.1f}%. "
                f"Одобрено только {approved} из {approved + rejected} партий."
            )
            q_rec = (
                "Необходимо немедленно приостановить отгрузку сомнительных партий, "
                "провести расследование причин и скорректировать технологический процесс."
            )
        self._word_section(doc, "3. ВЫПУСК ПО СМЕНАМ И КАЧЕСТВО ПРОДУКЦИИ",
                           shift_bullets, q_conclusion, q_rec,
                           chart_bytes=self._chart_po_quality_params(q_summary))

        # ── Section 4: Продажи по регионам ──────────────────────────────────
        region_list = getattr(regions, "regions", []) or []
        total_rev = sum(
            float(r.get("total_amount", 0) if isinstance(r, dict) else getattr(r, "total_amount", 0))
            for r in region_list
        )
        sales_bullets: list[str] = [f"• Общая выручка за период: {total_rev:,.0f} руб."]
        for r in region_list[:6]:
            reg = r.get("region", "—") if isinstance(r, dict) else getattr(r, "region", "—")
            amt = float(r.get("total_amount", 0) if isinstance(r, dict) else getattr(r, "total_amount", 0))
            pct = float(r.get("percentage", 0) if isinstance(r, dict) else getattr(r, "percentage", 0))
            sales_bullets.append(f"• {reg}: {amt:,.0f} руб. ({pct:.1f}%)")
        top_region = (
            (region_list[0].get("region") if isinstance(region_list[0], dict)
             else getattr(region_list[0], "region", "—"))
            if region_list else "—"
        )
        sales_conclusion = (
            f"Выручка за период составила {total_rev:,.0f} руб. "
            f"Ведущий регион — «{top_region}». "
            "Анализ региональной структуры позволяет оптимизировать логистику и приоритеты поставок."
        ) if region_list else "Данные по продажам за период отсутствуют."
        sales_rec = (
            "Рекомендуется сконцентрировать маркетинговые усилия на регионах с высоким потенциалом роста "
            "и обеспечить стабильность поставок в ключевые регионы."
        )
        self._word_section(doc, "4. ПРОДАЖИ ПО РЕГИОНАМ",
                           sales_bullets, sales_conclusion, sales_rec,
                           chart_bytes=self._chart_po_sales_pie(regions))

        # ── Section 5: Сенсоры и Запасы ─────────────────────────────────────
        sensor_items = sensors.get("items", []) or []
        total_alerts = sum(int(s.get("alert_count", 0) or 0) for s in sensor_items)
        total_readings = sum(int(s.get("reading_count", 0) or 0) for s in sensor_items)
        alert_ratio = (total_alerts / total_readings) if total_readings else 0.0
        _, _, sensor_level = self._assess_defect(Decimal(str(round(alert_ratio, 4))))

        sensor_bullets: list[str] = [
            f"• Параметров контролируется: {len(sensor_items)}",
            f"• Всего показаний: {total_readings}, тревог: {total_alerts} ({alert_ratio * 100:.1f}%)",
        ]
        for s in sensor_items[:5]:
            name = s.get("parameter_name", "—")
            alerts = int(s.get("alert_count", 0) or 0)
            readings = int(s.get("reading_count", 0) or 1)
            ratio = alerts / readings
            _, _, lv = self._assess_defect(Decimal(str(round(ratio, 4))))
            sensor_bullets.append(
                f"• {name}: среднее {float(s.get('avg_value') or 0):.2f} {s.get('unit', '')}, "
                f"тревог {alerts}/{readings} — {STATUS_LABEL[lv]}"
            )

        inv_items = inventory.get("items", []) or []
        snap_date = inventory.get("snapshot_date", "—")
        sensor_bullets.append(f"• Запасы (срез {snap_date}): {len(inv_items)} позиций на складе")

        if sensor_level == NORMAL:
            sensor_conclusion = "Производственное оборудование работает в штатном режиме. Уровень тревог в норме."
            sensor_rec = "Рекомендуется соблюдать регламент технического обслуживания для поддержания стабильности."
        elif sensor_level == WARNING:
            sensor_conclusion = (
                f"Зафиксировано повышенное количество тревог сенсоров: "
                f"{total_alerts} из {total_readings} показаний."
            )
            sensor_rec = "Рекомендуется провести профилактическое обслуживание оборудования и проверить датчики."
        else:
            sensor_conclusion = (
                f"Критический уровень тревог: {total_alerts} из {total_readings} показаний "
                f"({alert_ratio * 100:.1f}%). Возможны сбои оборудования."
            )
            sensor_rec = (
                "Необходимо срочно задействовать службы технического обслуживания, "
                "проверить исправность критически важных узлов и при необходимости остановить линию."
            )
        self._word_section(doc, "5. СОСТОЯНИЕ ОБОРУДОВАНИЯ И ЗАПАСЫ",
                           sensor_bullets, sensor_conclusion, sensor_rec)

        # ── Overall conclusion ────────────────────────────────────────────────
        overall = self._worst([worst_kpi, worst_orders, worst_q, sensor_level])
        doc.add_heading("ИТОГОВЫЙ ВЫВОД И РЕКОМЕНДАЦИИ", 1)
        if overall == NORMAL:
            doc.add_paragraph(
                "Производство функционирует в штатном режиме по всем ключевым направлениям. "
                "KPI, качество, выполнение заказов и состояние оборудования соответствуют плановым показателям. "
                "Рекомендуется продолжать плановую работу и мониторинг показателей."
            )
        elif overall == WARNING:
            doc.add_paragraph(
                "Выявлен ряд отклонений, требующих внимания руководства. "
                "Ситуация управляема, однако без корректирующих мер возможно ухудшение показателей. "
                "Рекомендуется инициировать план корректирующих мероприятий по проблемным направлениям."
            )
        else:
            doc.add_paragraph(
                "Зафиксированы критические отклонения по одному или нескольким направлениям. "
                "Требуется незамедлительное вмешательство и эскалация на уровень высшего руководства. "
                "Необходимо разработать план восстановления и установить контрольные точки."
            )

        return self._word_to_bytes(doc)
